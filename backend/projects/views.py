from django.db.models import Q
from django.contrib.auth import get_user_model
from rest_framework import viewsets, permissions
from rest_framework.exceptions import ValidationError, PermissionDenied
from .models import Project, ProjectMember
from .serializers import ProjectSerializer, ProjectMemberSerializer
from django.http import HttpResponse
from docx import Document
from rest_framework.decorators import action
from rest_framework.response import Response
from notifications.utils import notify_project
from notifications.models import Notification
from html.parser import HTMLParser

User = get_user_model()


class _DocxHtmlParser(HTMLParser):
    """Converts simple TipTap-generated HTML into paragraphs/runs on a python-docx Document."""

    def __init__(self, document):
        super().__init__()
        self.document = document
        self.paragraph = None
        self.bold = False
        self.italic = False
        self.list_type = None

    def _ensure_paragraph(self, style=None):
        self.paragraph = self.document.add_paragraph(style=style)

    def handle_starttag(self, tag, attrs):
        if tag == 'p':
            self._ensure_paragraph()
        elif tag in ('h1', 'h2', 'h3'):
            level = {'h1': 1, 'h2': 2, 'h3': 3}[tag]
            self.paragraph = self.document.add_heading('', level=level)
        elif tag in ('strong', 'b'):
            self.bold = True
        elif tag in ('em', 'i'):
            self.italic = True
        elif tag == 'li':
            self._ensure_paragraph(style='List Bullet')
        elif tag == 'br':
            if self.paragraph is not None:
                self.paragraph.add_run().add_break()

    def handle_endtag(self, tag):
        if tag in ('strong', 'b'):
            self.bold = False
        elif tag in ('em', 'i'):
            self.italic = False

    def handle_data(self, data):
        if not data.strip():
            return
        if self.paragraph is None:
            self._ensure_paragraph()
        run = self.paragraph.add_run(data)
        run.bold = self.bold
        run.italic = self.italic


def add_html_to_docx(document, html_content):
    if not html_content or not html_content.strip():
        document.add_paragraph('')
        return
    parser = _DocxHtmlParser(document)
    parser.feed(html_content)


class ProjectViewSet(viewsets.ModelViewSet):
    serializer_class = ProjectSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return Project.objects.filter(
            Q(owner=self.request.user) | Q(members__user=self.request.user)
        ).select_related('owner').distinct()

    def perform_create(self, serializer):
        serializer.save(owner=self.request.user)

    def perform_destroy(self, instance):
        if instance.owner != self.request.user:
            raise PermissionDenied('Faqat loyiha egasi uni o\'chira oladi.')
        instance.delete()

    @action(detail=True, methods=['get'])
    def export_docx(self, request, pk=None):
        project = self.get_object()
        chapters = project.chapters.all().order_by('order')

        document = Document()
        document.add_heading(project.title, level=0)
        if project.description:
            document.add_paragraph(project.description)

        for chapter in chapters:
            document.add_heading(chapter.title, level=1)
            add_html_to_docx(document, chapter.content)

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{project.title}.docx"'
        document.save(response)
        return response


class ProjectMemberViewSet(viewsets.ModelViewSet):
    serializer_class = ProjectMemberSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        queryset = ProjectMember.objects.filter(
            Q(project__owner=user) | Q(project__members__user=user)
        ).select_related('user', 'project').distinct()
        project_id = self.request.query_params.get('project')
        if project_id:
            queryset = queryset.filter(project_id=project_id)
        return queryset

    def perform_create(self, serializer):
        email = serializer.validated_data.pop('email')
        try:
            new_user = User.objects.get(email__iexact=email)
        except User.DoesNotExist:
            raise ValidationError({'email': 'Bu email bilan foydalanuvchi topilmadi'})
        member = serializer.save(user=new_user)
        actor = self.request.user
        actor_name = actor.first_name or actor.email
        notify_project(
            project=member.project,
            actor=new_user,
            verb='member_joined',
            message=f'{actor_name} yangi hamkorni loyihaga qo\'shdi',
            link=f'/projects/{member.project.id}',
        )
        Notification.objects.create(
            recipient=new_user,
            project=member.project,
            verb='member_joined',
            message=f'{actor_name} sizni "{member.project.title}" loyihasiga qo\'shdi',
            link=f'/projects/{member.project.id}',
        )